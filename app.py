import os
import json
import re
import uuid
import pickle
import numpy as np
import requests
from flask import Flask, render_template, request, jsonify, flash, redirect, url_for, send_file, session
from flask_sqlalchemy import SQLAlchemy
from flask_admin import Admin
from flask_admin.contrib.sqla import ModelView
from werkzeug.utils import secure_filename
from datetime import datetime
from io import BytesIO
from docx import Document
from groq import Groq
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup

# ================================================================
# APP CONFIGURATION
# ================================================================
app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///nigat.db'
app.config['SECRET_KEY'] = 'mysecretkey'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# FAISS storage directory
FAISS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'faiss_indexes')
if not os.path.exists(FAISS_DIR):
    os.makedirs(FAISS_DIR)

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# ================================================================
# DATABASE INITIALIZATION
# ================================================================
db = SQLAlchemy(app)

# ================================================================
# GROQ API CONFIGURATION
# ================================================================

GROQ_API_KEY = os.environ.get('GROQ_API_KEY', '')
GROQ_MODEL = "llama-3.3-70b-versatile"
FALLBACK_MODELS = ["mixtral-8x7b-32768", "gemma2-9b-it", "llama-3.1-8b-instant"]

groq_client = None
if GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        print("✅ Groq client initialized successfully")
    except Exception as e:
        print(f"⚠️ Failed to initialize Groq client: {e}")
else:
    print("⚠️ GROQ_API_KEY not set. Please set it in environment variables.")

# ================================================================
# WEB SEARCH FUNCTIONS (DuckDuckGo + Google fallback)
# ================================================================

def web_search(query, max_results=5):
    """Search the web using DuckDuckGo; fallback to Google if needed"""
    try:
        with DDGS() as ddgs:
            results = []
            for r in ddgs.text(query, max_results=max_results):
                results.append({
                    'title': r.get('title', ''),
                    'body': r.get('body', ''),
                    'href': r.get('href', '')
                })
            if results:
                return results
    except Exception as e:
        print(f"⚠️ DuckDuckGo search error: {e}")
    
    # Fallback: Google search using requests and BeautifulSoup
    try:
        return google_search(query, max_results)
    except Exception as e:
        print(f"⚠️ Google search fallback error: {e}")
        return []

def google_search(query, max_results=5):
    """Simple Google search using requests and BeautifulSoup (fallback)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    url = f"https://www.google.com/search?q={requests.utils.quote(query)}&num={max_results}"
    response = requests.get(url, headers=headers, timeout=10)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, 'html.parser')
    results = []
    for g in soup.find_all('div', class_='g'):
        title_elem = g.find('h3')
        if not title_elem:
            continue
        title = title_elem.get_text()
        link_elem = g.find('a')
        if not link_elem:
            continue
        href = link_elem.get('href')
        if not href:
            continue
        # Extract snippet
        snippet_elem = g.find('div', class_='VwiC3b')
        snippet = snippet_elem.get_text() if snippet_elem else ''
        results.append({
            'title': title,
            'body': snippet,
            'href': href
        })
        if len(results) >= max_results:
            break
    return results

def format_search_results(results):
    """Format search results for context"""
    if not results:
        return "No web search results available."
    formatted = []
    for i, r in enumerate(results, 1):
        formatted.append(f"[{i}] {r['title']}\n{r['body']}\nSource: {r['href']}\n")
    return "\n".join(formatted)

# ================================================================
# PAGE RANGE EXTRACTION FROM USER QUERY
# ================================================================

def extract_page_range(query):
    """Extract page range like 'from page 10 to 20' or 'pages 5-12'"""
    pattern = r'(?:from\s+)?pages?\s*([0-9]+)\s*(?:-|to|–)\s*([0-9]+)'
    match = re.search(pattern, query, re.IGNORECASE)
    if match:
        start = int(match.group(1))
        end = int(match.group(2))
        if start <= end:
            return start, end
    # Also try single page
    single_pattern = r'page\s*([0-9]+)'
    match = re.search(single_pattern, query, re.IGNORECASE)
    if match:
        p = int(match.group(1))
        return p, p
    return None, None

# ================================================================
# AI RESPONSE FUNCTION (Groq + Web Search + Fallback)
# ================================================================

def get_ai_response(system_prompt, user_query, context_chunks=None, use_web_search=False, page_range=None):
    """Get AI response using Groq with optional web search and fallback models"""
    
    if not groq_client:
        return "⚠️ Groq API key is not set. Please add GROQ_API_KEY to environment variables."
    
    # Detect language
    if detect_language(user_query) == 'amharic':
        lang_instruction = "You MUST respond in Amharic (በአማርኛ)."
    else:
        lang_instruction = "You MUST respond in English."
    
    # Build context
    context_text = ""
    
    # Add document context if available (filter by page range if specified)
    if context_chunks:
        context_text += "\n=== DOCUMENT CONTEXT ===\n"
        # If page range is specified, we try to include only those chunks (if we have page numbers)
        # For simplicity, we just include all chunks and let the model focus on the range if mentioned
        context_text += "\n\n---\n\n".join(context_chunks[:5])
    
    # Add web search results if requested
    if use_web_search:
        # If page range is specified, we can refine the search query
        search_query = user_query
        if page_range:
            start, end = page_range
            search_query += f" pages {start} to {end}"
        search_results = web_search(search_query)
        if search_results:
            context_text += "\n=== WEB SEARCH RESULTS ===\n"
            context_text += format_search_results(search_results)
    
    if not context_text.strip():
        context_text = "No context available. Please upload a document or enable web search for specific information."
    
    # Detect prompt type
    prompt_type = detect_prompt_type(user_query)
    prompt_template = get_prompt_template(prompt_type, page_range)
    
    # Build full prompt
    full_prompt = (
        f"{system_prompt}\n\n"
        f"=== LANGUAGE ===\n{lang_instruction}\n\n"
        f"=== TASK ===\n{prompt_template}\n\n"
        f"=== CONTEXT ===\n{context_text}\n\n"
        f"=== USER QUESTION ===\n{user_query}"
    )
    
    # Try primary model first
    models_to_try = [GROQ_MODEL] + FALLBACK_MODELS
    
    for model in models_to_try:
        try:
            print(f"🤖 Trying model: {model}")
            completion = groq_client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": full_prompt}
                ],
                temperature=0.1,
                max_tokens=2048,
                top_p=0.95
            )
            response = completion.choices[0].message.content
            print(f"✅ Success with model: {model}")
            return response
        except Exception as e:
            error_msg = str(e)
            print(f"⚠️ Model {model} failed: {error_msg[:100]}")
            if "decommissioned" in error_msg:
                print(f"⏭️ Model {model} is decommissioned, trying next...")
                continue
            continue
    
    return "⚠️ All available models failed. Please try again later or check your Groq API key."

# ================================================================
# PROMPT MANAGEMENT
# ================================================================

def detect_prompt_type(query):
    query_lower = query.lower()
    if any(w in query_lower for w in ['daily lesson', 'ዕለታዊ', 'lesson plan']):
        return 'daily_lesson'
    elif any(w in query_lower for w in ['annual', 'ዓመታዊ', 'yearly']):
        return 'annual_plan'
    elif any(w in query_lower for w in ['semester', 'ሴሚስተር']):
        return 'semester_plan'
    elif any(w in query_lower for w in ['monthly', 'ወርሃዊ']):
        return 'monthly_plan'
    elif any(w in query_lower for w in ['weekly', 'ሳምንታዊ']):
        return 'weekly_plan'
    elif any(w in query_lower for w in ['exam', 'test', 'quiz', 'ፈተና', 'ምዘና']):
        return 'exam'
    elif any(w in query_lower for w in ['summary', 'summarize', 'ማጠቃለያ']):
        return 'summary'
    return 'general'

def get_prompt_template(prompt_type, page_range=None):
    page_info = ""
    if page_range:
        start, end = page_range
        page_info = f" (focus on content from pages {start} to {end})"
    
    base_templates = {
        'daily_lesson': f"""
=== DETAILED DAILY LESSON PLAN {page_info} ===
Generate a COMPLETE and DETAILED daily lesson plan with this structure.
Use information from the provided context (document or web search) to fill in the content.
If the context contains specific information about the topic, use it directly.

STRUCTURE:

1. SCHOOL INFORMATION: School Name, Teacher Name, Grade/Section, Subject, Date, Unit, Topic, Page
2. LESSON OBJECTIVES: At least 3-5 specific, measurable objectives (Bloom's Taxonomy)
3. RATIONALE: Why this lesson is important, connection to previous/future lessons
4. PREREQUISITES: What students should already know
5. COMPETENCIES: 3-5 specific competencies students will demonstrate
6. LESSON STAGES: A DETAILED TABLE with: Stage | Time | Teacher Activities | Student Activities | Methodology | Assessment
7. TEACHING AIDS: List all materials, visual aids, technology needed
8. SUPPORT FOR LEARNERS: Table with: Category (Slow/Medium/Fast) | Support Strategies
9. ASSESSMENT: Formative and Summative
10. APPROVALS: Table with: Role | Name | Signature | Date
11. TEACHER'S SELF-ASSESSMENT: What went well, what could be improved
""",
        'annual_plan': f"""
=== ANNUAL LESSON PLAN {page_info} ===
Generate an annual lesson plan with this structure:
1. SCHOOL INFORMATION
2. TABLE: Month | Week | Topics | Objectives | Methodology | Evaluation
""",
        'semester_plan': f"""
=== SEMESTER LESSON PLAN {page_info} ===
1. SCHOOL INFORMATION
2. TABLE: Week | Topic | Objectives | Activities | Assessment
""",
        'monthly_plan': f"""
=== MONTHLY LESSON PLAN {page_info} ===
1. SCHOOL INFORMATION
2. TABLE: Week | Topic | Objectives | Activities | Resources
""",
        'weekly_plan': f"""
=== WEEKLY LESSON PLAN {page_info} ===
1. SCHOOL INFORMATION
2. TABLE: Day | Topic | Objectives | Activities | Homework
""",
        'exam': f"""
=== EXAM GENERATOR {page_info} ===
Generate exam questions based on the provided content.
Include: Multiple choice (≥10), True/False (≥5), Short answer (≥5), Essay (≥2).
""",
        'summary': f"""
=== SUMMARY GENERATOR {page_info} ===
Create a comprehensive summary: main topics, key points, important concepts, key terms, study tips.
""",
        'general': f"""
=== GENERAL ASSISTANCE {page_info} ===
Answer the user's question based on the provided context. Be helpful, accurate, and cite specific information.
"""
    }
    return base_templates.get(prompt_type, base_templates['general'])

# ================================================================
# HELPER FUNCTIONS
# ================================================================

def detect_language(text):
    if not text:
        return 'english'
    amharic_pattern = re.compile(r'[\u1200-\u137F]')
    if amharic_pattern.search(text):
        return 'amharic'
    return 'english'

def remove_duplicate_sentences(text):
    if not text:
        return text
    sentences = re.split(r'(?<=[.!?])\s+', text)
    seen = set()
    unique = []
    for sentence in sentences:
        s = sentence.strip()
        if not s or len(s) < 5:
            continue
        norm = s.lower()
        if norm not in seen:
            seen.add(norm)
            unique.append(s)
    return ' '.join(unique)

# ================================================================
# MEMORY OPTIMIZED PDF EXTRACTION (50 Pages)
# ================================================================

def extract_pdf_text_streaming(filepath):
    """Extract PDF text page by page with memory optimization (up to 50 pages)"""
    try:
        import pdfplumber
        text_parts = []
        total_pages = 0
        page_texts = []  # store each page text with page number
        
        with pdfplumber.open(filepath) as pdf:
            total_pages = len(pdf.pages)
            max_pages = min(total_pages, 50)
            
            for i in range(max_pages):
                page = pdf.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    page_texts.append((i+1, page_text))  # store page number and text
                page = None
                if (i + 1) % 10 == 0:
                    print(f"📄 Extracted page {i+1}/{max_pages}")
        
        full_text = "\n\n".join(text_parts)
        text_parts = None
        return full_text, max_pages if full_text else "No text found in PDF.", 0, page_texts
    except Exception as e:
        return f"PDF extraction error: {str(e)}", 0, []

# ================================================================
# EMBEDDING MODEL (Singleton)
# ================================================================

_embedding_model = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            print("🔄 Loading embedding model...")
            _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            print("✅ Embedding model loaded")
        except ImportError:
            print("⚠️ sentence-transformers not installed!")
            _embedding_model = None
    return _embedding_model

def get_embedding(text):
    model = get_embedding_model()
    if model is None:
        import hashlib
        words = text.lower().split()
        vector = np.zeros(384)
        for word in words[:100]:
            h = hashlib.md5(word.encode()).hexdigest()
            for i in range(min(8, len(h))):
                vector[i % 384] += (int(h[i], 16) - 8) / 16
        norm = np.linalg.norm(vector)
        return vector / (norm + 1e-8)
    try:
        return model.encode(text, normalize_embeddings=True)
    except:
        return np.zeros(384)

# ================================================================
# MEMORY OPTIMIZED RAG SYSTEM (with page numbers)
# ================================================================

class EnterpriseRAG:
    def __init__(self):
        self.doc_metadata = {}
        self.chunk_texts = {}
        self.chunk_pages = {}  # store page number for each chunk
        self.faiss_indexes = {}
        self.chunk_size = 300
        self.overlap = 50
        self.max_chunks = 400
    
    def get_index_path(self, session_id):
        return os.path.join(FAISS_DIR, f"{session_id}.faiss")
    
    def get_metadata_path(self, session_id):
        return os.path.join(FAISS_DIR, f"{session_id}_meta.pkl")
    
    def _chunk_text_with_pages(self, text, page_texts):
        """Split text into chunks and assign page numbers"""
        # page_texts is list of (page_num, page_text)
        # We need to split text into chunks and map to page numbers
        # Simple approach: concatenate page texts and then split; track which page each chunk belongs to
        chunks = []
        chunk_pages = []
        current_page = 1
        current_text = ""
        for page_num, page_text in page_texts:
            # Append page text with a marker
            if current_text:
                current_text += "\n\n"
            current_text += f"[PAGE {page_num}]\n" + page_text
            # Split current_text into chunks if it exceeds chunk_size
            words = current_text.split()
            while len(words) > self.chunk_size:
                chunk_words = words[:self.chunk_size]
                chunk = ' '.join(chunk_words)
                chunks.append(chunk)
                # Determine page number for this chunk: find the last [PAGE X] marker
                # Extract page numbers from chunk
                pages_in_chunk = re.findall(r'\[PAGE (\d+)\]', chunk)
                if pages_in_chunk:
                    # Use the last page number in chunk
                    chunk_pages.append(int(pages_in_chunk[-1]))
                else:
                    chunk_pages.append(current_page)
                words = words[self.chunk_size - self.overlap:]
                current_text = ' '.join(words)
            # After processing, keep remaining text for next page
            current_text = ' '.join(words)
            current_page = page_num
        # Final chunks
        if current_text.strip():
            chunks.append(current_text)
            pages_in_chunk = re.findall(r'\[PAGE (\d+)\]', current_text)
            if pages_in_chunk:
                chunk_pages.append(int(pages_in_chunk[-1]))
            else:
                chunk_pages.append(current_page)
        return chunks, chunk_pages
    
    def store_document(self, session_id, text, filename, pages=0, page_texts=None):
        self.doc_metadata[session_id] = {
            'filename': filename,
            'pages': pages,
            'word_count': len(text.split()),
            'chunk_count': 0
        }
        
        try:
            import faiss
            embedding_dim = 384
            faiss_index = faiss.IndexFlatIP(embedding_dim)
        except ImportError:
            faiss_index = None
        
        chunks = []
        chunk_pages = []
        if page_texts:
            chunks, chunk_pages = self._chunk_text_with_pages(text, page_texts)
        else:
            # Fallback: chunk without page info
            for chunk in self._chunk_text_streaming(text):
                chunks.append(chunk)
                chunk_pages.append(None)  # unknown page
        
        # Store chunks and pages
        self.chunk_texts[session_id] = chunks
        self.chunk_pages[session_id] = chunk_pages
        self.doc_metadata[session_id]['chunk_count'] = len(chunks)
        
        # Build FAISS index
        embeddings = []
        for chunk in chunks:
            emb = get_embedding(chunk)
            embeddings.append(emb)
            if faiss_index is not None and len(embeddings) >= 30:
                emb_array = np.array(embeddings).astype('float32')
                faiss_index.add(emb_array)
                embeddings = []
                import gc; gc.collect()
        if embeddings and faiss_index is not None:
            emb_array = np.array(embeddings).astype('float32')
            faiss_index.add(emb_array)
        
        if faiss_index is not None:
            faiss_path = self.get_index_path(session_id)
            faiss.write_index(faiss_index, faiss_path)
            self.faiss_indexes[session_id] = faiss_index
        
        # Save metadata with page info
        meta_path = self.get_metadata_path(session_id)
        with open(meta_path, 'wb') as f:
            pickle.dump({
                'chunks': chunks,
                'chunk_pages': chunk_pages,
                'metadata': self.doc_metadata[session_id]
            }, f)
        
        print(f"📚 Stored {len(chunks)} chunks with page info")
        return len(chunks)
    
    def _chunk_text_streaming(self, text):
        words = text.split()
        total_words = len(words)
        chunk_count = 0
        for start in range(0, total_words, self.chunk_size - self.overlap):
            if chunk_count >= self.max_chunks:
                break
            end = min(start + self.chunk_size, total_words)
            chunk_words = words[start:end]
            if len(chunk_words) < 10:
                continue
            chunk_count += 1
            yield ' '.join(chunk_words)
            if end >= total_words:
                break
    
    def _load_metadata(self, session_id):
        meta_path = self.get_metadata_path(session_id)
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'rb') as f:
                    data = pickle.load(f)
                    self.chunk_texts[session_id] = data.get('chunks', [])
                    self.chunk_pages[session_id] = data.get('chunk_pages', [])
                    self.doc_metadata[session_id] = data.get('metadata', {})
                    return data
            except:
                pass
        return None
    
    def _load_faiss_index(self, session_id):
        if session_id in self.faiss_indexes:
            return self.faiss_indexes[session_id]
        try:
            import faiss
            faiss_path = self.get_index_path(session_id)
            if os.path.exists(faiss_path):
                index = faiss.read_index(faiss_path)
                self.faiss_indexes[session_id] = index
                return index
        except:
            pass
        return None
    
    def get_relevant_chunks(self, session_id, query, max_tokens=4000, page_range=None):
        if session_id not in self.chunk_texts:
            self._load_metadata(session_id)
        
        if session_id not in self.chunk_texts or not self.chunk_texts[session_id]:
            return []
        
        chunks = self.chunk_texts[session_id]
        chunk_pages = self.chunk_pages.get(session_id, [])
        
        # If page range is specified, filter chunks by page
        if page_range:
            start_page, end_page = page_range
            filtered = []
            for i, chunk in enumerate(chunks):
                if i < len(chunk_pages) and chunk_pages[i] is not None:
                    if start_page <= chunk_pages[i] <= end_page:
                        filtered.append(chunk)
                else:
                    # If page unknown, include it (fallback)
                    filtered.append(chunk)
            if filtered:
                chunks = filtered
        
        faiss_index = self._load_faiss_index(session_id)
        
        if faiss_index is not None:
            try:
                import faiss
                query_emb = get_embedding(query)
                query_emb = np.array([query_emb]).astype('float32')
                
                k = min(15, len(chunks))
                scores, indices = faiss_index.search(query_emb, k)
                
                selected = []
                total_tokens = 0
                for idx in indices[0]:
                    if idx < 0 or idx >= len(chunks):
                        continue
                    chunk = chunks[idx]
                    estimated = len(chunk) // 4
                    if total_tokens + estimated <= max_tokens:
                        selected.append(chunk)
                        total_tokens += estimated
                    if len(selected) >= 5:
                        break
                return selected if selected else [chunks[0]]
            except:
                pass
        
        # Keyword fallback
        query_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', query.lower()))
        scored = []
        for i, chunk in enumerate(chunks[:100]):
            chunk_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', chunk.lower()))
            overlap = len(query_words & chunk_words)
            if overlap > 0:
                scored.append((overlap, chunk))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [chunk for _, chunk in scored[:4]]
    
    def get_document_info(self, session_id):
        if session_id in self.doc_metadata:
            return self.doc_metadata[session_id]
        self._load_metadata(session_id)
        return self.doc_metadata.get(session_id)
    
    def clear(self, session_id):
        if session_id in self.doc_metadata:
            del self.doc_metadata[session_id]
        if session_id in self.chunk_texts:
            del self.chunk_texts[session_id]
        if session_id in self.chunk_pages:
            del self.chunk_pages[session_id]
        if session_id in self.faiss_indexes:
            del self.faiss_indexes[session_id]
        
        faiss_path = self.get_index_path(session_id)
        if os.path.exists(faiss_path):
            os.remove(faiss_path)
        meta_path = self.get_metadata_path(session_id)
        if os.path.exists(meta_path):
            os.remove(meta_path)

rag = EnterpriseRAG()

# ================================================================
# MODELS (same as before)
# ================================================================
class Course(db.Model):
    __tablename__ = 'course'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=False)
    color = db.Column(db.String(20), default='blue')
    quiz_link = db.Column(db.String(500))

class AnnualPlan(db.Model):
    __tablename__ = 'annual_plan'
    id = db.Column(db.Integer, primary_key=True)
    school_name = db.Column(db.String(200))
    teacher_name = db.Column(db.String(100))
    subject = db.Column(db.String(100))
    grade = db.Column(db.Integer)
    section = db.Column(db.String(20))
    year = db.Column(db.Integer)
    total_days = db.Column(db.Integer)
    unit_number = db.Column(db.Integer)
    unit_title = db.Column(db.String(200))
    unit_objectives = db.Column(db.Text)
    month_data = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class LaboratoryPlan(db.Model):
    __tablename__ = 'laboratory_plan'
    id = db.Column(db.Integer, primary_key=True)
    school_name = db.Column(db.String(200))
    teacher_name = db.Column(db.String(100))
    subject = db.Column(db.String(100))
    grade = db.Column(db.Integer)
    year = db.Column(db.Integer)
    experiment_data = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class DailyPlan(db.Model):
    __tablename__ = 'daily_plan'
    id = db.Column(db.Integer, primary_key=True)
    teacher_name = db.Column(db.String(100))
    school_name = db.Column(db.String(200))
    grade = db.Column(db.Integer)
    section = db.Column(db.String(20))
    subject = db.Column(db.String(100))
    date = db.Column(db.Date)
    unit_number = db.Column(db.Integer)
    lesson_topic = db.Column(db.String(200))
    page = db.Column(db.String(20))
    rationale = db.Column(db.Text)
    prerequisites = db.Column(db.Text)
    competencies = db.Column(db.Text)
    starter_time = db.Column(db.Integer)
    starter_teacher = db.Column(db.Text)
    starter_student = db.Column(db.Text)
    starter_method = db.Column(db.String(100))
    starter_assessment = db.Column(db.String(100))
    starter_aids = db.Column(db.String(200))
    main_time = db.Column(db.Integer)
    main_teacher = db.Column(db.Text)
    main_student = db.Column(db.Text)
    main_method = db.Column(db.String(100))
    main_assessment = db.Column(db.String(100))
    main_aids = db.Column(db.String(200))
    conclude_time = db.Column(db.Integer)
    conclude_teacher = db.Column(db.Text)
    conclude_student = db.Column(db.Text)
    conclude_method = db.Column(db.String(100))
    conclude_assessment = db.Column(db.String(100))
    conclude_aids = db.Column(db.String(200))
    slow_learners = db.Column(db.Text)
    medium_learners = db.Column(db.Text)
    fast_learners = db.Column(db.Text)
    self_assessment = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class PeaceClubPlan(db.Model):
    __tablename__ = 'peace_club_plan'
    id = db.Column(db.Integer, primary_key=True)
    school_name = db.Column(db.String(200))
    district = db.Column(db.String(200))
    woreda = db.Column(db.String(100))
    school_level = db.Column(db.String(200))
    club_name = db.Column(db.String(200), default='Peace Club / የሰላም ክበብ')
    teacher_name = db.Column(db.String(100))
    teacher_signature = db.Column(db.String(100))
    secretary_name = db.Column(db.String(100))
    secretary_signature = db.Column(db.String(100))
    year = db.Column(db.Integer)
    month = db.Column(db.String(20))
    vision = db.Column(db.Text)
    mission = db.Column(db.Text)
    opportunities = db.Column(db.Text)
    challenges = db.Column(db.Text)
    solutions = db.Column(db.Text)
    action_plan = db.Column(db.Text)
    student_members = db.Column(db.Text)
    teacher_members = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def get_action_plan(self):
        return json.loads(self.action_plan) if self.action_plan else []
    
    def get_student_members(self):
        return json.loads(self.student_members) if self.student_members else []
    
    def get_teacher_members(self):
        return json.loads(self.teacher_members) if self.teacher_members else []

class PeaceClubActivity(db.Model):
    __tablename__ = 'peace_club_activity'
    id = db.Column(db.Integer, primary_key=True)
    club_plan_id = db.Column(db.Integer, db.ForeignKey('peace_club_plan.id'))
    activity_number = db.Column(db.Integer)
    activity_name = db.Column(db.String(500))
    hamle = db.Column(db.Boolean, default=False)
    nehase = db.Column(db.Boolean, default=False)
    meskerem = db.Column(db.Boolean, default=False)
    tikimt = db.Column(db.Boolean, default=False)
    hidar = db.Column(db.Boolean, default=False)
    tahsas = db.Column(db.Boolean, default=False)
    tir = db.Column(db.Boolean, default=False)
    yekatit = db.Column(db.Boolean, default=False)
    megabit = db.Column(db.Boolean, default=False)
    miazia = db.Column(db.Boolean, default=False)
    ginbot = db.Column(db.Boolean, default=False)
    sene = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

admin = Admin(app, name='Nigat Admin')
admin.add_view(ModelView(Course, db))
admin.add_view(ModelView(AnnualPlan, db))
admin.add_view(ModelView(LaboratoryPlan, db))
admin.add_view(ModelView(DailyPlan, db))
admin.add_view(ModelView(PeaceClubPlan, db))
admin.add_view(ModelView(PeaceClubActivity, db))

# ================================================================
# ROUTES (same as before with minor changes)
# ================================================================
@app.route('/')
def home():
    try:
        courses = Course.query.all()
        return render_template('index.html', courses=courses)
    except Exception as e:
        print(f"❌ Home route error: {e}")
        with app.app_context():
            db.create_all()
        courses = Course.query.all()
        return render_template('index.html', courses=courses)

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        flash('Registration submitted successfully!', 'success')
        return redirect(url_for('home'))
    return render_template('register.html')

@app.route('/course/<int:course_id>')
def course_detail(course_id):
    return render_template('course_detail.html', course=Course.query.get_or_404(course_id))

@app.route('/upload_text', methods=['POST'])
def upload_text():
    """Upload text directly (up to 10 pages)"""
    data = request.json
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'success': False, 'message': 'No text provided.'}), 400
    word_count = len(text.split())
    pages = (word_count // 250) + 1
    if pages > 10:
        return jsonify({'success': False, 'message': f'Text exceeds 10 pages ({pages} pages). Please reduce the content.'}), 400
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
    session_id = session['session_id']
    filename = f"text_input_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    # For text, we don't have page numbers, so pass None for page_texts
    num_chunks = rag.store_document(session_id, text, filename, pages, page_texts=None)
    session['text_filename'] = filename
    session['text_pages'] = pages
    session['text_chunks'] = num_chunks
    return jsonify({
        'success': True,
        'message': f'Text uploaded and indexed! ({pages} pages, {num_chunks} chunks)',
        'session_id': session_id,
        'pages': pages,
        'chunks': num_chunks
    }), 200

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected.'}), 400
    if file and file.filename.lower().endswith('.pdf'):
        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            file_size = os.path.getsize(filepath) / (1024 * 1024)
            text, pages, page_texts = extract_pdf_text_streaming(filepath)
            if not text or text.startswith("PDF extraction error"):
                return jsonify({'success': False, 'message': f'Error extracting text: {text}'}), 500
            if pages > 50:
                return jsonify({'success': False, 'message': f'PDF has {pages} pages. Maximum is 50 pages.'}), 400
            if 'session_id' not in session:
                session['session_id'] = str(uuid.uuid4())
            session_id = session['session_id']
            num_chunks = rag.store_document(session_id, text, filename, pages, page_texts)
            session['pdf_filename'] = filename
            session['pdf_size'] = file_size
            session['pdf_pages'] = pages
            session['pdf_chunks'] = num_chunks
            return jsonify({
                'success': True, 
                'message': f'PDF uploaded and indexed! ({file_size:.1f}MB, {pages} pages, {num_chunks} chunks)',
                'session_id': session_id,
                'pages': pages,
                'chunks': num_chunks
            }), 200
        except Exception as e:
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
    else:
        return jsonify({'success': False, 'message': 'Only PDF files are allowed.'}), 400

@app.route('/upload_image', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return jsonify({'error': 'No image uploaded'}), 400
    file = request.files['image']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'bmp'}
    if '.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        return jsonify({'message': 'Image uploaded successfully'}), 200
    else:
        return jsonify({'error': 'Unsupported file type'}), 400

@app.route('/clear_context', methods=['POST'])
def clear_context():
    session_id = session.get('session_id')
    if session_id:
        rag.clear(session_id)
    session.pop('pdf_filename', None)
    session.pop('pdf_size', None)
    session.pop('pdf_pages', None)
    session.pop('pdf_chunks', None)
    session.pop('text_filename', None)
    session.pop('text_pages', None)
    session.pop('text_chunks', None)
    return jsonify({'message': 'Context cleared successfully'}), 200

# ================================================================
# AI CHAT ROUTE (With Web Search and Page Range)
# ================================================================

@app.route('/ask_ai', methods=['POST'])
def ask_ai():
    user_query = request.json.get('query', '').strip()
    use_web_search = request.json.get('use_web_search', False)
    
    if not user_query:
        return jsonify({"answer": "Please ask a question."})
    
    query_lang = detect_language(user_query)
    print(f"🔍 Detected language: {query_lang}")
    print(f"📝 User query: {user_query[:100]}...")
    
    # Extract page range from query
    start_page, end_page = extract_page_range(user_query)
    page_range = (start_page, end_page) if start_page is not None else None
    if page_range:
        print(f"📄 Page range detected: {start_page} to {end_page}")
    
    session_id = session.get('session_id')
    relevant_chunks = []
    doc_info = None
    
    if session_id:
        relevant_chunks = rag.get_relevant_chunks(session_id, user_query, max_tokens=4000, page_range=page_range)
        doc_info = rag.get_document_info(session_id)
        if doc_info:
            print(f"📄 Document: {doc_info.get('filename', 'unknown')} ({doc_info.get('pages', 0)} pages, {doc_info.get('chunk_count', 0)} chunks)")
    
    print(f"📚 Retrieved {len(relevant_chunks)} relevant chunks")
    
    # Auto-enable web search if no context and it's a lesson plan or general query
    is_lesson_plan = any(w in user_query.lower() for w in ['lesson plan', 'daily lesson', 'annual plan', 'semester', 'monthly', 'weekly', 'daily plan'])
    if not relevant_chunks and not use_web_search:
        use_web_search = True
        print("🌐 Auto-enabling web search")
    
    if query_lang == 'amharic':
        language_instruction = "You MUST respond in Amharic (በአማርኛ)."
    else:
        language_instruction = "You MUST respond in English."
    
    system_prompt = (
        "You are 'Nigat AI Tutor'. Created by Teacher Fisaha Melke.\n\n"
        f"=== LANGUAGE RULE ===\n{language_instruction}\n"
        "Do NOT switch languages.\n\n"
        "=== ABOUT THE CREATOR ===\n"
        "My name is Fisiha Melke. I graduated from Ambo University with a Bachelor's degree in Biology in 2024 (2016 E.C.). I have more than two years of teaching experience in private schools. I hold a Certificate in Video Editing. I am currently developing Nigat Tutor AI, an educational platform designed to support Ethiopian teachers and students.\n\n"
        "=== ACCURACY RULE ===\n"
        "Provide ONLY accurate information based on the provided context. If the context doesn't contain the answer, use web search results. If neither contains the answer, say so clearly.\n\n"
        "=== AMHARIC SPELLING ===\n"
        "Correct spellings: 'ጎንደር' (not ንንደር/ጀንደር), 'ኢትዮጵያ' (not እትዮጵያ).\n\n"
        "=== WEB SEARCH ===\n"
        f"Web search is {'ENABLED' if use_web_search else 'DISABLED'} for this query.\n"
    )
    
    answer = get_ai_response(
        system_prompt, 
        user_query, 
        relevant_chunks if relevant_chunks else None,
        use_web_search=use_web_search,
        page_range=page_range
    )
    
    if answer:
        answer = remove_duplicate_sentences(answer)
    
    return jsonify({
        "answer": answer or "⚠️ No AI response available. Please try again later.",
        "used_web_search": use_web_search
    })

# ================================================================
# DOWNLOAD WORD (same as before)
# ================================================================

@app.route('/download_word', methods=['POST'])
def download_word():
    data = request.json
    content = data.get('content', '')
    filename = data.get('filename', 'Nigat_AI_Response')
    if not content:
        return jsonify({'error': 'No content to download'}), 400
    try:
        doc = Document()
        doc.add_heading('Nigat AI Tutor Response', 0)
        lines = content.split('\n')
        in_table = False
        table_rows = []
        table_headers = []
        for line in lines:
            line = line.strip()
            if line.startswith('|') and line.endswith('|'):
                cells = [cell.strip() for cell in line[1:-1].split('|')]
                if all('---' in cell or ':' in cell for cell in cells):
                    continue
                if not in_table:
                    in_table = True
                    table_headers = cells
                else:
                    table_rows.append(cells)
            else:
                if in_table and table_rows:
                    num_cols = max(len(table_headers), max([len(row) for row in table_rows]) if table_rows else 0)
                    if num_cols > 0 and table_headers:
                        table = doc.add_table(rows=1 + len(table_rows), cols=num_cols)
                        table.style = 'Table Grid'
                        for i, header in enumerate(table_headers[:num_cols]):
                            cell = table.cell(0, i)
                            cell.text = header
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                        for row_idx, row in enumerate(table_rows):
                            for col_idx, cell_text in enumerate(row[:num_cols]):
                                table.cell(row_idx + 1, col_idx).text = cell_text
                    table_rows = []
                    table_headers = []
                    in_table = False
                if line:
                    if line.startswith('#'):
                        heading_level = min(len(line) - len(line.lstrip('#')), 6)
                        heading_text = line.lstrip('#').strip()
                        doc.add_heading(heading_text, level=heading_level)
                    else:
                        doc.add_paragraph(line)
        if in_table and table_rows:
            num_cols = max(len(table_headers), max([len(row) for row in table_rows]) if table_rows else 0)
            if num_cols > 0 and table_headers:
                table = doc.add_table(rows=1 + len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                for i, header in enumerate(table_headers[:num_cols]):
                    cell = table.cell(0, i)
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                for row_idx, row in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row[:num_cols]):
                        table.cell(row_idx + 1, col_idx).text = cell_text
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        safe_filename = re.sub(r'[^\w\s-]', '', filename)
        safe_filename = re.sub(r'[-\s]+', '_', safe_filename)
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{safe_filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"❌ Word download error: {e}")
        return jsonify({'error': f'Failed to generate document: {str(e)}'}), 500

# ================================================================
# LESSON PLAN ROUTES (unchanged)
# ================================================================

@app.route('/lesson')
def lesson_home():
    annual_plans = AnnualPlan.query.all()
    laboratory_plans = LaboratoryPlan.query.all()
    daily_plans = DailyPlan.query.all()
    return render_template('lesson_plan.html', 
                         annual_plans=annual_plans,
                         laboratory_plans=laboratory_plans,
                         daily_plans=daily_plans)

@app.route('/lesson/annual', methods=['GET', 'POST'])
def annual_plan():
    if request.method == 'POST':
        months = request.form.getlist('month[]')
        weeks = request.form.getlist('week[]')
        periods = request.form.getlist('period[]')
        date_ranges = request.form.getlist('date_range[]')
        pages = request.form.getlist('page[]')
        topics = request.form.getlist('topics[]')
        objectives = request.form.getlist('objectives[]')
        methodologies = request.form.getlist('methodology[]')
        teaching_aids = request.form.getlist('teaching_aids[]')
        evaluations = request.form.getlist('evaluation[]')
        month_data = []
        for i in range(len(months)):
            month_data.append({
                'month': months[i],
                'week': weeks[i],
                'period': periods[i],
                'date_range': date_ranges[i],
                'page': pages[i],
                'topics': topics[i],
                'objectives': objectives[i],
                'methodology': methodologies[i],
                'teaching_aids': teaching_aids[i],
                'evaluation': evaluations[i]
            })
        plan = AnnualPlan(
            school_name=request.form.get('school_name'),
            teacher_name=request.form.get('teacher_name'),
            subject=request.form.get('subject'),
            grade=int(request.form.get('grade')) if request.form.get('grade') else None,
            section=request.form.get('section'),
            year=int(request.form.get('year')) if request.form.get('year') else None,
            total_days=int(request.form.get('total_days')) if request.form.get('total_days') else None,
            unit_number=int(request.form.get('unit_number')) if request.form.get('unit_number') else None,
            unit_title=request.form.get('unit_title'),
            unit_objectives=request.form.get('unit_objectives'),
            month_data=str(month_data)
        )
        db.session.add(plan)
        db.session.commit()
        flash('Annual plan created successfully!', 'success')
        return redirect(url_for('lesson_home'))
    return render_template('annual_plan_form.html')

@app.route('/lesson/laboratory', methods=['GET', 'POST'])
def laboratory_plan():
    if request.method == 'POST':
        exp_numbers = request.form.getlist('exp_number[]')
        exp_titles = request.form.getlist('exp_title[]')
        apparatus_list = request.form.getlist('apparatus[]')
        chemicals_list = request.form.getlist('chemicals[]')
        unit_numbers = request.form.getlist('unit_number[]')
        pages = request.form.getlist('page[]')
        months = request.form.getlist('month[]')
        dates = request.form.getlist('date[]')
        experiment_data = []
        for i in range(len(exp_numbers)):
            experiment_data.append({
                'exp_number': exp_numbers[i],
                'exp_title': exp_titles[i],
                'apparatus': apparatus_list[i],
                'chemicals': chemicals_list[i],
                'unit_number': unit_numbers[i],
                'page': pages[i],
                'month': months[i],
                'date': dates[i]
            })
        plan = LaboratoryPlan(
            school_name=request.form.get('school_name'),
            teacher_name=request.form.get('teacher_name'),
            subject=request.form.get('subject'),
            grade=int(request.form.get('grade')) if request.form.get('grade') else None,
            year=int(request.form.get('year')) if request.form.get('year') else None,
            experiment_data=str(experiment_data)
        )
        db.session.add(plan)
        db.session.commit()
        flash('Laboratory plan created successfully!', 'success')
        return redirect(url_for('lesson_home'))
    return render_template('laboratory_plan_form.html')

@app.route('/lesson/daily', methods=['GET', 'POST'])
def daily_plan():
    if request.method == 'POST':
        date_str = request.form.get('date')
        plan = DailyPlan(
            teacher_name=request.form.get('teacher_name'),
            school_name=request.form.get('school_name'),
            grade=int(request.form.get('grade')) if request.form.get('grade') else None,
            section=request.form.get('section'),
            subject=request.form.get('subject'),
            date=datetime.strptime(date_str, '%Y-%m-%d').date() if date_str else None,
            unit_number=int(request.form.get('unit_number')) if request.form.get('unit_number') else None,
            lesson_topic=request.form.get('lesson_topic'),
            page=request.form.get('page'),
            rationale=request.form.get('rationale'),
            prerequisites=request.form.get('prerequisites'),
            competencies=request.form.get('competencies'),
            starter_time=int(request.form.get('starter_time')) if request.form.get('starter_time') else None,
            starter_teacher=request.form.get('starter_teacher'),
            starter_student=request.form.get('starter_student'),
            starter_method=request.form.get('starter_method'),
            starter_assessment=request.form.get('starter_assessment'),
            starter_aids=request.form.get('starter_aids'),
            main_time=int(request.form.get('main_time')) if request.form.get('main_time') else None,
            main_teacher=request.form.get('main_teacher'),
            main_student=request.form.get('main_student'),
            main_method=request.form.get('main_method'),
            main_assessment=request.form.get('main_assessment'),
            main_aids=request.form.get('main_aids'),
            conclude_time=int(request.form.get('conclude_time')) if request.form.get('conclude_time') else None,
            conclude_teacher=request.form.get('conclude_teacher'),
            conclude_student=request.form.get('conclude_student'),
            conclude_method=request.form.get('conclude_method'),
            conclude_assessment=request.form.get('conclude_assessment'),
            conclude_aids=request.form.get('conclude_aids'),
            slow_learners=request.form.get('slow_learners'),
            medium_learners=request.form.get('medium_learners'),
            fast_learners=request.form.get('fast_learners'),
            self_assessment=request.form.get('self_assessment')
        )
        db.session.add(plan)
        db.session.commit()
        flash('Daily plan created successfully!', 'success')
        return redirect(url_for('lesson_home'))
    return render_template('daily_plan_form.html')

# ================================================================
# PEACE CLUB ROUTES (unchanged)
# ================================================================

@app.route('/peaceclub')
def peaceclub_home():
    club_plans = PeaceClubPlan.query.all()
    return render_template('peaceclub_home.html', club_plans=club_plans)

@app.route('/peaceclub/create', methods=['GET', 'POST'])
def peaceclub_create():
    if request.method == 'POST':
        plan = PeaceClubPlan(
            school_name=request.form.get('school_name'),
            district=request.form.get('district'),
            woreda=request.form.get('woreda'),
            school_level=request.form.get('school_level'),
            club_name=request.form.get('club_name', 'Peace Club / የሰላም ክበብ'),
            teacher_name=request.form.get('teacher_name'),
            teacher_signature=request.form.get('teacher_signature'),
            secretary_name=request.form.get('secretary_name'),
            secretary_signature=request.form.get('secretary_signature'),
            year=int(request.form.get('year')) if request.form.get('year') else None,
            month=request.form.get('month'),
            vision=request.form.get('vision'),
            mission=request.form.get('mission'),
            opportunities=request.form.get('opportunities'),
            challenges=request.form.get('challenges'),
            solutions=request.form.get('solutions')
        )
        db.session.add(plan)
        db.session.flush()
        activity_names = request.form.getlist('activity_name[]')
        hamle_values = request.form.getlist('hamle')
        nehase_values = request.form.getlist('nehase')
        meskerem_values = request.form.getlist('meskerem')
        tikimt_values = request.form.getlist('tikimt')
        hidar_values = request.form.getlist('hidar')
        tahsas_values = request.form.getlist('tahsas')
        tir_values = request.form.getlist('tir')
        yekatit_values = request.form.getlist('yekatit')
        megabit_values = request.form.getlist('megabit')
        miazia_values = request.form.getlist('miazia')
        ginbot_values = request.form.getlist('ginbot')
        sene_values = request.form.getlist('sene')
        for i, name in enumerate(activity_names):
            if name.strip():
                activity = PeaceClubActivity(
                    club_plan_id=plan.id,
                    activity_number=i + 1,
                    activity_name=name.strip(),
                    hamle=str(i) in hamle_values,
                    nehase=str(i) in nehase_values,
                    meskerem=str(i) in meskerem_values,
                    tikimt=str(i) in tikimt_values,
                    hidar=str(i) in hidar_values,
                    tahsas=str(i) in tahsas_values,
                    tir=str(i) in tir_values,
                    yekatit=str(i) in yekatit_values,
                    megabit=str(i) in megabit_values,
                    miazia=str(i) in miazia_values,
                    ginbot=str(i) in ginbot_values,
                    sene=str(i) in sene_values
                )
                db.session.add(activity)
        student_names = request.form.getlist('student_name[]')
        student_grades = request.form.getlist('student_grade[]')
        student_data = []
        for i in range(len(student_names)):
            if student_names[i].strip():
                student_data.append({
                    'name': student_names[i].strip(),
                    'grade': student_grades[i] if i < len(student_grades) else ''
                })
        plan.student_members = json.dumps(student_data)
        teacher_names = request.form.getlist('teacher_name[]')
        teacher_grades = request.form.getlist('teacher_grade[]')
        teacher_data = []
        for i in range(len(teacher_names)):
            if teacher_names[i].strip():
                teacher_data.append({
                    'name': teacher_names[i].strip(),
                    'grade': teacher_grades[i] if i < len(teacher_grades) else ''
                })
        plan.teacher_members = json.dumps(teacher_data)
        db.session.commit()
        flash('Peace Club plan created successfully!', 'success')
        return redirect(url_for('peaceclub_home'))
    return render_template('peaceclub_create.html')

@app.route('/peaceclub/view/<int:plan_id>')
def peaceclub_view(plan_id):
    plan = PeaceClubPlan.query.get_or_404(plan_id)
    activities = PeaceClubActivity.query.filter_by(club_plan_id=plan_id).order_by(PeaceClubActivity.activity_number).all()
    student_members = plan.get_student_members()
    teacher_members = plan.get_teacher_members()
    return render_template('peaceclub_view.html', 
                         plan=plan, 
                         activities=activities,
                         student_members=student_members,
                         teacher_members=teacher_members)

@app.route('/peaceclub/edit/<int:plan_id>', methods=['GET', 'POST'])
def peaceclub_edit(plan_id):
    plan = PeaceClubPlan.query.get_or_404(plan_id)
    activities = PeaceClubActivity.query.filter_by(club_plan_id=plan_id).order_by(PeaceClubActivity.activity_number).all()
    student_members = plan.get_student_members()
    teacher_members = plan.get_teacher_members()
    if request.method == 'POST':
        plan.school_name = request.form.get('school_name')
        plan.district = request.form.get('district')
        plan.woreda = request.form.get('woreda')
        plan.school_level = request.form.get('school_level')
        plan.club_name = request.form.get('club_name', 'Peace Club / የሰላም ክበብ')
        plan.teacher_name = request.form.get('teacher_name')
        plan.teacher_signature = request.form.get('teacher_signature')
        plan.secretary_name = request.form.get('secretary_name')
        plan.secretary_signature = request.form.get('secretary_signature')
        plan.year = int(request.form.get('year')) if request.form.get('year') else None
        plan.month = request.form.get('month')
        plan.vision = request.form.get('vision')
        plan.mission = request.form.get('mission')
        plan.opportunities = request.form.get('opportunities')
        plan.challenges = request.form.get('challenges')
        plan.solutions = request.form.get('solutions')
        for activity in activities:
            db.session.delete(activity)
        activity_names = request.form.getlist('activity_name[]')
        hamle_values = request.form.getlist('hamle')
        nehase_values = request.form.getlist('nehase')
        meskerem_values = request.form.getlist('meskerem')
        tikimt_values = request.form.getlist('tikimt')
        hidar_values = request.form.getlist('hidar')
        tahsas_values = request.form.getlist('tahsas')
        tir_values = request.form.getlist('tir')
        yekatit_values = request.form.getlist('yekatit')
        megabit_values = request.form.getlist('megabit')
        miazia_values = request.form.getlist('miazia')
        ginbot_values = request.form.getlist('ginbot')
        sene_values = request.form.getlist('sene')
        for i, name in enumerate(activity_names):
            if name.strip():
                activity = PeaceClubActivity(
                    club_plan_id=plan.id,
                    activity_number=i + 1,
                    activity_name=name.strip(),
                    hamle=str(i) in hamle_values,
                    nehase=str(i) in nehase_values,
                    meskerem=str(i) in meskerem_values,
                    tikimt=str(i) in tikimt_values,
                    hidar=str(i) in hidar_values,
                    tahsas=str(i) in tahsas_values,
                    tir=str(i) in tir_values,
                    yekatit=str(i) in yekatit_values,
                    megabit=str(i) in megabit_values,
                    miazia=str(i) in miazia_values,
                    ginbot=str(i) in ginbot_values,
                    sene=str(i) in sene_values
                )
                db.session.add(activity)
        student_names = request.form.getlist('student_name[]')
        student_grades = request.form.getlist('student_grade[]')
        student_data = []
        for i in range(len(student_names)):
            if student_names[i].strip():
                student_data.append({
                    'name': student_names[i].strip(),
                    'grade': student_grades[i] if i < len(student_grades) else ''
                })
        plan.student_members = json.dumps(student_data)
        teacher_names = request.form.getlist('teacher_name[]')
        teacher_grades = request.form.getlist('teacher_grade[]')
        teacher_data = []
        for i in range(len(teacher_names)):
            if teacher_names[i].strip():
                teacher_data.append({
                    'name': teacher_names[i].strip(),
                    'grade': teacher_grades[i] if i < len(teacher_grades) else ''
                })
        plan.teacher_members = json.dumps(teacher_data)
        db.session.commit()
        flash('Peace Club plan updated successfully!', 'success')
        return redirect(url_for('peaceclub_view', plan_id=plan.id))
    return render_template('peaceclub_edit.html', 
                         plan=plan, 
                         activities=activities,
                         student_members=student_members,
                         teacher_members=teacher_members)

@app.route('/peaceclub/delete/<int:plan_id>')
def peaceclub_delete(plan_id):
    plan = PeaceClubPlan.query.get_or_404(plan_id)
    activities = PeaceClubActivity.query.filter_by(club_plan_id=plan_id).all()
    for activity in activities:
        db.session.delete(activity)
    db.session.delete(plan)
    db.session.commit()
    flash('Peace Club plan deleted successfully!', 'success')
    return redirect(url_for('peaceclub_home'))

# ================================================================
# CREATE TABLES ON APPLICATION STARTUP
# ================================================================
with app.app_context():
    try:
        db.create_all()
        print("✅ Database tables created/verified successfully.")
        print(f"📊 Using database: {app.config['SQLALCHEMY_DATABASE_URI']}")
        print(f"🤖 Groq API: {'✅ Configured' if GROQ_API_KEY else '❌ Not configured'}")
        print(f"🧠 Groq Model: {GROQ_MODEL}")
        print(f"📄 Max PDF pages: 50")
        print(f"📝 Max text pages: 10")
        print(f"🌐 Web search: ✅ Enabled (auto-enabled for lesson plans)")
        print(f"📖 Page range support: ✅ Enabled")
    except Exception as e:
        print(f"❌ Failed to create tables: {e}")

# ================================================================
# MAIN
# ================================================================
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
